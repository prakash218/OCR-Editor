import time
start = time.time()
from tkinter import *
from tkinter import messagebox
from PIL import ImageTk, Image
from tkinter import filedialog
import threading
import pytesseract
import cv2
import numpy as np
from shutil import copy
from tkinter.filedialog import asksaveasfile
from docx2pdf import convert
from os import remove,system
import docx
import unicodedata

pytesseract.pytesseract.tesseract_cmd = r'Tesseract-OCR\tesseract.exe'
root = Tk()


#Constants
from constants import *
filename = "start.png"



#Variables
black_and_white = IntVar()
threshold = IntVar()
img = Image.open("start.png")
img = img.resize((500,281) ,  Image.ANTIALIAS)
img = ImageTk.PhotoImage(img)
save_path = None
export_path = None
output = None
text = None
 


#functions

def remove_control_characters(s):
	"""this function is to remove non ASCII characters

	Args:
		s (str): string

	Returns:
		str: string with only ASCII characters
	"""
	res = ''
	for i in s:
		try:
			if unicodedata.category(i)[0] != 'C' or i == '\n':
				res += i
		except:
			print(i)
	return res

def get_file(path):
	"""To export the file as docx or pdf file

	Args:
		path (str): [path of the file to be exported]
	"""
	global text
	extension = path.split('.')[1]
	filename = path.split('.')[0].split('/')[-1]
	print(filename)
	doc = docx.Document()
	doc.add_heading('Output', 0) 
	text = remove_control_characters(text)
	
	text_copy = text.split('\n')
	final = []

	for line in text_copy:
		if line.isspace() or line == '':
			continue
		final.append(line)


	if extension != 'pdf':
		for line in final:
			para = doc.add_paragraph(line)
			run = para.add_run()
			run.add_break()
		doc.save(path)
	else:
		filename = filename + '.docx'
		try:
			remove(filename)
		except:
			pass
		for line in final:
			para = doc.add_paragraph(line)
			run = para.add_run()
			run.add_break()
		doc.save(filename)
		convert(filename,path)
		remove(filename)




def export():
	"""to export a file which is already saved
	"""
	global export_path
	if export_path == None:
		export_as()
	else:
		get_file(export_path)

def export_as():
	"""To export the file as pdf or docx
	"""
	global export_path
	files = [("Text files","*.docx"),
			 ("PDF files","*.pdf"),
			 ("all files","*.*")] 
	try:
		export_path = asksaveasfile(filetypes = files, defaultextension = files).name 
	except:
		return
	
	get_file(export_path)

def save():
	"""To save the edited image
	"""
	global save_path
	if save_path == None:
		save_as()
	else:
		copy(edit_name,save_path)

def save_as():
	"""To edit the image as png or jpeg
	"""
	global save_path,edit_name
	files = [("png files","*.png"),
			 ("jpg files","*.jpg"),
			 ("all files","*.*")] 
	try:
		save_path = asksaveasfile(filetypes = files, defaultextension = files).name 
	except:
		return
	copy(edit_name,save_path)

def open_edit():
	"""To open a new image
	"""
	image = cv2.imread(edit_name)
	cv2.imshow("Edited", image)


def change_image(path):
	"""To change the preview of the edited image

	Args:
		path ([str]): [the path of the edited image]
	"""
	global img
	img = Image.open(path)
	img = img.resize((500,281) ,  Image.ANTIALIAS)
	img = ImageTk.PhotoImage(img)
	original.configure(image = img)

def refresh():
	"""Refreshing the edited image --- works on a seperate thread
	"""
	global edit_name,text
	black = 0
	curr = None
	thresh = 0
	med_blur = 1
	should_change = 0
	t_start = 0
	t_end = 0
	alphanum = 'abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ1234567890!@#$%^&*()-_=+`~"\',.<>/?{}[]\\|'

	while True:
		global image,filename
		prev_file = curr
		curr = filename
		
		image = cv2.imread(filename)

		prev_black = black
		black = black_and_white.get()

		prev_thresh = thresh
		thresh = threshold.get()

		prev_med_blur = med_blur
		med_blur = median_blur.get()

		prev_t_start = t_start
		prev_t_end = t_end

		t_start = thresh_start.get()
		t_end = thresh_end.get()

		gray = image

		if black:
			gray = cv2.cvtColor(gray, cv2.COLOR_BGR2GRAY)

		if thresh:
			gray = cv2.threshold(gray,t_start,t_end,cv2.THRESH_BINARY)[1]

		if thresh and (t_start != prev_t_start or t_end != prev_t_end):
			should_change = 1
		else:
			should_change = 0

		# if prev_med_blur != med_blur:
		if not med_blur % 2:
			med_blur += 1
		gray = cv2.medianBlur(gray,med_blur)


		

		edit_name = 'edit.png'
		
		# print(pytesseract.image_to_boxes(Image.open(edit_name)))


		if prev_black != black or prev_file != curr or prev_thresh != thresh or prev_med_blur != med_blur or should_change:

			cv2.imwrite(edit_name, gray)
			
			img = Image.open(edit_name)
			img = img.resize((500,281) ,  Image.ANTIALIAS)
			img = ImageTk.PhotoImage(img)
			edit.configure(image = img)
			text = pytesseract.image_to_string(Image.open(edit_name))
			for i in alphanum:
				if i in text:
					output.delete('1.0',END)
					output.insert('1.0',text)
					break
			else:
				output.delete('1.0',END)



def open_file():
    """To open a photo"""
    global image
    filename_copy = filedialog.askopenfilename(initialdir = "/",title = "Select a File",filetypes = (("image files","*.png*"),
    																								 ("image files","*.jpg*"),
    																								 ("all files","*.*"))) 
    
    # Change label contents 
    if filename_copy:
    	global filename
    	filename = filename_copy
    	change_image(filename)
	    
def about():
	system('start documentation.pdf')


thread = threading.Thread(target = refresh, daemon = True)
thread.start()


"""
root : screen object for tkinter
"""
root.title('OCR')
root.geometry('1500x800+200+100')
root.configure(bg = BG)



menubar = Menu(root)  
file = Menu(menubar, tearoff = 0)
file.add_command(label="Open", command = open_file)  
file.add_command(label="Save",command = save)  
file.add_command(label="Save as...",command = save_as) 
file.add_command(label="Export",command = export)
file.add_command(label="Export as...",command = export_as) 
file.add_command(label="Close")  
  
file.add_separator()  
  
file.add_command(label="Exit", command=root.quit)  
  
menubar.add_cascade(label="File", menu=file)  
 
  
 
help = Menu(menubar, tearoff=0)  
help.add_command(label="About", command = about)  
menubar.add_cascade(label="Help", menu=help)  

frame1 = Frame(root, bd = 5, relief = 'ridge', bg = FG)
frame1.place(x = 650, y = 40, height = 750, width = 300)

# frame2 = Frame(root, bd = 5, relief = 'ridge', bg = FG)
# frame2.place(x = 1200, y = 40, height = 750, width = 300)


root.config(menu = menubar)
Label(root, text = 'Original', font = ('Helvetica',36), bg = BG, fg = SECONDARY).grid(row = 1, column = 1) 
Label(root, text = 'Edited', font = ('Helvetica',36), bg = BG, fg = SECONDARY).grid(row = 3, column = 1) 

original = Label(root, image = img, bg= BG)
original.grid(row = 2 , column = 1,padx = 30, pady = 30)

edit = Label(root, image = img, bg= BG)
edit.grid(row = 4, column = 1, pady = 30)

Checkbutton(frame1, text = "Black and White:", font = ('Helvetica',20), bg = FG, fg = SECONDARY,activebackground = SECONDARY, activeforeground = BG,variable = black_and_white).grid(row = 1,sticky='w')
Checkbutton(frame1, text = "Threshold:", font = ('Helvetica',20), bg = FG, fg = SECONDARY,activebackground = SECONDARY, activeforeground = FG,variable = threshold).grid(row = 2,sticky='w')


Label(frame1, text = "Threshold Start", font = ('Helvetica',20), bg = FG, fg = SECONDARY).grid(row = 3,sticky='nsew')
thresh_start = Scale(frame1, from_ = 1, to_ = 255, orient = HORIZONTAL, bg = FG, fg = SECONDARY)
thresh_start.grid(row = 4,column = 0,sticky='w',ipadx = 80,padx = 10)

Label(frame1, text = "Threshold End", font = ('Helvetica',20), bg = FG, fg = SECONDARY).grid(row = 5,sticky='nsew')
thresh_end = Scale(frame1, from_ = 1, to_ = 255, orient = HORIZONTAL, bg = FG, fg = SECONDARY)
thresh_end.grid(row = 6,column = 0,sticky='w',ipadx = 80,padx = 10)

Label(frame1, text = "Median Blur:", font = ('Helvetica',20), bg = FG, fg = SECONDARY).grid(row = 7)
median_blur = Scale(frame1, from_ = 1, to_ = 49, orient = HORIZONTAL, bg = FG, fg = SECONDARY)
median_blur.grid(row = 8,column = 0,sticky='w',ipadx = 80,padx = 10)

Label(root, text = 'Output', font = ('Helvetica',20), bg = BG, fg = SECONDARY).place(x = 1200, y = 10)
output = Text(root, font = ('Helvetica',12), relief = 'ridge', bd = 5)
output.place(x = 1000, y = 45, width = 480, height = 740)

Button(frame1, text = 'open edited', font = ('Helvetica',16),bg = FG, fg = SECONDARY, activebackground = SECONDARY, activeforeground = FG,command = open_edit).grid(row = 25, pady = 25)



end = time.time()
print(end - start)

root.mainloop()
